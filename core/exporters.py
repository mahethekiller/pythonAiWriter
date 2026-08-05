"""
Document exporters and output format handlers.
"""

import re
import html
from pathlib import Path
from typing import Union
from bs4 import BeautifulSoup
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_slug(text: str) -> str:
    """Generates a clean URL slug from title string."""
    text = text.lower().strip()
    text = re.sub(r'[^\w\s-]', '', text)
    text = re.sub(r'[\s_-]+', '-', text)
    return text.strip('-') or "article"


class DocxExporter:
    """Utility class to convert HTML strings to styled Word (.docx) documents."""

    @staticmethod
    def html_to_docx(html_content: str, output_path: Union[str, Path]) -> Path:
        """Parses HTML body content and exports a formatted .docx file."""
        doc = docx.Document()
        
        # Set standard margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        soup = BeautifulSoup(html_content, "html.parser")
        
        # Target body or main container, falling back to full soup
        body = soup.find("body") or soup.find("main") or soup

        # Extract block-level tags
        block_tags = body.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "ul", "ol", "blockquote"])

        if not block_tags:
            # Fallback: Extract non-empty text lines into paragraphs
            raw_text = body.get_text()
            lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
            for line in lines:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(6)
                p.add_run(line)
        else:
            for element in block_tags:
                if element.name in ["ul", "ol"]:
                    style = 'List Bullet' if element.name == "ul" else 'List Number'
                    for li in element.find_all("li", recursive=False):
                        p = doc.add_paragraph(style=style)
                        p.paragraph_format.space_after = Pt(3)
                        DocxExporter._add_html_runs(p, li)
                elif element.name == "li":
                    # Handled recursively by parent ul/ol
                    continue
                elif element.name == "h1":
                    p = doc.add_heading(level=1)
                    run = p.add_run(element.get_text())
                    run.font.size = Pt(20)
                    run.font.color.rgb = RGBColor(15, 23, 42)
                    run.font.bold = True
                elif element.name == "h2":
                    p = doc.add_heading(level=2)
                    run = p.add_run(element.get_text())
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor(30, 58, 138)
                    run.font.bold = True
                elif element.name == "h3":
                    p = doc.add_heading(level=3)
                    run = p.add_run(element.get_text())
                    run.font.size = Pt(13)
                    run.font.color.rgb = RGBColor(37, 99, 235)
                    run.font.bold = True
                elif element.name in ["h4", "h5", "h6"]:
                    p = doc.add_heading(level=4)
                    run = p.add_run(element.get_text())
                    run.font.size = Pt(11.5)
                    run.font.bold = True
                elif element.name in ["p", "blockquote"]:
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(6)
                    DocxExporter._add_html_runs(p, element)

        out_file = Path(output_path)
        doc.save(str(out_file))
        return out_file

    @staticmethod
    def _add_html_runs(paragraph, parent_element):
        """Recursively parses child nodes to add formatted runs (bold, italic, links)."""
        for child in parent_element.contents:
            if isinstance(child, str):
                if child.strip():
                    paragraph.add_run(child)
            elif child.name in ["strong", "b"]:
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name in ["em", "i"]:
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == "a":
                href = child.get("href", "")
                text = child.get_text()
                run = paragraph.add_run(f"{text} ({href})" if href else text)
                run.font.color.rgb = RGBColor(37, 99, 235)
                run.underline = True
            elif child.name == "code":
                run = paragraph.add_run(child.get_text())
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
            else:
                DocxExporter._add_html_runs(paragraph, child)
